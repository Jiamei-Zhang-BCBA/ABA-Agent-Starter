# api/tests/test_file_processor.py
"""
Unit tests for api/app/services/file_processor.py — parse_file()

TDD coverage:
  - .txt / .md  → UTF-8 happy path and GBK fallback decoding
  - .docx       → paragraph extraction via python-docx
  - .pdf        → page extraction via PyPDF2 (real file) and corrupted-PDF error path
  - .jpg / .png → P3 OCR placeholder
  - .mp3 / .wav → P3 Whisper placeholder
  - unknown ext → unsupported-type message

No external services are touched; all fixtures create in-memory bytes.
"""

from __future__ import annotations

import io
import struct

import pytest

from app.services.file_processor import parse_file


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _minimal_pdf_bytes() -> bytes:
    """
    Build the smallest valid single-page PDF that PyPDF2 can read.
    The page carries no visible text, but the file must parse without error.
    We do this without reportlab to avoid an extra dependency.
    """
    # A hand-crafted minimal PDF (revision 1.0, empty page, no fonts)
    body = (
        b"%PDF-1.0\n"
        b"1 0 obj<</Type /Catalog /Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type /Pages /Kids[3 0 R] /Count 1>>endobj\n"
        b"3 0 obj<</Type /Page /MediaBox[0 0 3 3]>>endobj\n"
    )
    xref_offset = len(body)
    body += (
        b"xref\n"
        b"0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4 /Root 1 0 R>>\n"
        b"startxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF"
    )
    return body


def _docx_bytes(text: str) -> bytes:
    """Return in-memory .docx bytes containing a single paragraph."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# .txt and .md — encoding tests
# ---------------------------------------------------------------------------


class TestPlainTextFiles:
    def test_parse_txt_utf8(self):
        """UTF-8 encoded .txt content is returned as-is."""
        content = "Hello, ABA Supervisor! 你好"
        result = parse_file(content.encode("utf-8"), "session_note.txt")
        assert result == content

    def test_parse_md_file(self):
        """Markdown files are treated as plain text and returned verbatim."""
        content = "# Progress Note\n\nTarget: **mand training**\n\nTrials: 10/10"
        result = parse_file(content.encode("utf-8"), "note.md")
        assert result == content

    def test_parse_txt_gbk(self):
        """GBK-encoded Chinese text is decoded correctly as a UTF-8 fallback."""
        chinese = "行为分析师督导报告"
        gbk_bytes = chinese.encode("gbk")
        # Confirm the bytes are NOT valid UTF-8
        with pytest.raises(UnicodeDecodeError):
            gbk_bytes.decode("utf-8")
        result = parse_file(gbk_bytes, "report.txt")
        assert result == chinese

    def test_parse_txt_preserves_newlines(self):
        """Multi-line text content preserves line structure."""
        content = "Line one\nLine two\nLine three"
        result = parse_file(content.encode("utf-8"), "multi.txt")
        assert "Line one" in result
        assert "Line two" in result
        assert "Line three" in result

    def test_parse_txt_empty_file(self):
        """Empty .txt file returns an empty string without error."""
        result = parse_file(b"", "empty.txt")
        assert result == ""

    def test_parse_md_with_special_characters(self):
        """Markdown with Unicode and emojis is handled without error."""
        content = "## 个案进展 🎯\n\n- 强化物偏好 ✓\n- 辅助层级: 零延迟辅助"
        result = parse_file(content.encode("utf-8"), "progress.md")
        assert "个案进展" in result
        assert "辅助层级" in result


# ---------------------------------------------------------------------------
# .jpg / .jpeg / .png — image placeholder
# ---------------------------------------------------------------------------


class TestImagePlaceholder:
    def test_parse_jpg_placeholder(self):
        """JPEG file returns the P3 OCR placeholder referencing the filename."""
        result = parse_file(b"\xff\xd8\xff", "photo.jpg")
        assert "photo.jpg" in result
        assert "OCR" in result
        assert "P3" in result

    def test_parse_jpeg_extension(self):
        """.jpeg extension is treated the same as .jpg."""
        result = parse_file(b"\xff\xd8\xff", "scan.jpeg")
        assert "scan.jpeg" in result
        assert "OCR" in result

    def test_parse_png_placeholder(self):
        """PNG file returns the P3 OCR placeholder."""
        result = parse_file(b"\x89PNG", "chart.png")
        assert "chart.png" in result
        assert "P3" in result

    def test_image_placeholder_exact_format(self):
        """Image placeholder matches the documented format string exactly."""
        filename = "abc.jpg"
        expected = f"[图片文件: {filename} — OCR 功能将在 P3 阶段启用]"
        result = parse_file(b"\xff\xd8\xff", filename)
        assert result == expected


# ---------------------------------------------------------------------------
# .mp3 / .m4a / .wav — audio placeholder
# ---------------------------------------------------------------------------


class TestAudioPlaceholder:
    def test_parse_mp3_placeholder(self):
        """MP3 file returns the P3 Whisper placeholder referencing the filename."""
        result = parse_file(b"\xff\xfb", "session_audio.mp3")
        assert "session_audio.mp3" in result
        assert "P3" in result

    def test_parse_m4a_extension(self):
        """.m4a extension is treated as an audio placeholder."""
        result = parse_file(b"", "interview.m4a")
        assert "interview.m4a" in result
        assert "语音转写" in result

    def test_parse_wav_extension(self):
        """.wav extension is treated as an audio placeholder."""
        result = parse_file(b"RIFF", "recording.wav")
        assert "recording.wav" in result
        assert "P3" in result

    def test_audio_placeholder_exact_format(self):
        """Audio placeholder matches the documented format string exactly."""
        filename = "note.mp3"
        expected = f"[音频文件: {filename} — 语音转写功能将在 P3 阶段启用]"
        result = parse_file(b"\xff\xfb", filename)
        assert result == expected


# ---------------------------------------------------------------------------
# Unknown / unsupported extension
# ---------------------------------------------------------------------------


class TestUnknownExtension:
    def test_parse_unknown_extension(self):
        """.xyz returns the unsupported-type message containing the extension."""
        result = parse_file(b"some bytes", "data.xyz")
        assert ".xyz" in result
        assert "不支持" in result

    def test_parse_no_extension(self):
        """A filename with no extension is treated as unsupported."""
        result = parse_file(b"raw", "no_ext_file")
        # Path('no_ext_file').suffix == '' so ext is ''
        assert "不支持" in result

    def test_parse_csv_unsupported(self):
        """.csv is not in the supported set and returns the unsupported message."""
        result = parse_file(b"a,b,c\n1,2,3", "data.csv")
        assert ".csv" in result

    def test_unsupported_exact_format(self):
        """Unsupported message matches the documented format string exactly."""
        result = parse_file(b"", "file.xyz")
        assert result == "[不支持的文件类型: .xyz]"

    def test_extension_is_case_insensitive(self):
        """Upper-case extensions are normalised before lookup."""
        # .TXT in upper-case should be parsed as text, not unsupported
        content = "case test"
        result = parse_file(content.encode("utf-8"), "NOTE.TXT")
        assert result == content

    def test_jpeg_upper_case(self):
        """.JPEG upper-case returns the image placeholder, not unsupported."""
        result = parse_file(b"\xff\xd8\xff", "PHOTO.JPEG")
        assert "OCR" in result
        assert "不支持" not in result


# ---------------------------------------------------------------------------
# .pdf — PyPDF2 integration
# ---------------------------------------------------------------------------


class TestPdfParsing:
    def test_parse_pdf_minimal_valid(self):
        """A minimal valid PDF (empty page) parses without raising an exception."""
        pdf_bytes = _minimal_pdf_bytes()
        result = parse_file(pdf_bytes, "empty.pdf")
        # Empty page yields no text — result is either an empty string or
        # a page-separator-only string.  Either way it must NOT be an error message.
        assert "[PDF 解析失败" not in result

    def test_parse_pdf_corrupted_returns_error(self):
        """A corrupted / non-PDF byte sequence returns the error placeholder."""
        junk = b"this is definitely not a PDF"
        result = parse_file(junk, "corrupt.pdf")
        assert result.startswith("[PDF 解析失败:")

    def test_parse_pdf_empty_bytes(self):
        """Zero-byte PDF returns the error placeholder gracefully."""
        result = parse_file(b"", "empty.pdf")
        assert result.startswith("[PDF 解析失败:")

    def test_parse_pdf_with_pypdf2_writer(self):
        """PDF created by PdfWriter round-trips through parse_file without error."""
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        pdf_bytes = buf.read()

        result = parse_file(pdf_bytes, "blank.pdf")
        assert "[PDF 解析失败" not in result


# ---------------------------------------------------------------------------
# .docx — python-docx integration
# ---------------------------------------------------------------------------


class TestDocxParsing:
    def test_parse_docx_single_paragraph(self):
        """A .docx with one paragraph returns that paragraph's text."""
        text = "测试内容"
        result = parse_file(_docx_bytes(text), "note.docx")
        assert text in result

    def test_parse_docx_multiple_paragraphs(self):
        """Multiple paragraphs are all present in the output."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段")
        doc.add_paragraph("第三段")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parse_file(buf.read(), "multi.docx")
        assert "第一段" in result
        assert "第二段" in result
        assert "第三段" in result

    def test_parse_docx_empty_paragraphs_stripped(self):
        """Empty/whitespace-only paragraphs are not included in the output."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Real content")
        doc.add_paragraph("   ")  # whitespace-only — should be stripped
        doc.add_paragraph("")     # empty — should be stripped
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parse_file(buf.read(), "sparse.docx")
        assert "Real content" in result
        # The whitespace-only paragraph must not inflate output
        assert result.strip() == "Real content"

    def test_parse_docx_corrupted_returns_error(self):
        """A corrupted byte sequence returns the DOCX error placeholder."""
        junk = b"not a docx file"
        result = parse_file(junk, "corrupt.docx")
        assert result.startswith("[DOCX 解析失败:")

    def test_parse_docx_empty_bytes(self):
        """Zero-byte DOCX returns the error placeholder gracefully."""
        result = parse_file(b"", "empty.docx")
        assert result.startswith("[DOCX 解析失败:")

    def test_parse_docx_with_ascii_and_unicode(self):
        """DOCX containing mixed ASCII and Unicode text is extracted correctly."""
        text = "Behavior plan — 行为干预计划 (P1)"
        result = parse_file(_docx_bytes(text), "plan.docx")
        assert "Behavior plan" in result
        assert "行为干预计划" in result
